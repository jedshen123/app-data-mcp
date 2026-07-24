from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/lute/code/app-data-mcp")
SOURCE = ROOT / "docs" / "app-data-mcp技术实现方案-Q3.md"
OUT = ROOT / "docs" / "app-data-mcp技术实现方案-Q3.docx"
IMAGE_DIR = ROOT / "docs" / "images" / "q3-tech-plan"

FONT_REGULAR = "/System/Library/Fonts/STHeiti Light.ttc"
FONT_BOLD = "/System/Library/Fonts/STHeiti Medium.ttc"

NAVY = "123B6D"
BLUE = "1D67B1"
CYAN = "0A91B8"
GREEN = "17864B"
ORANGE = "D66A1F"
RED = "B42318"
INK = "172033"
MUTED = "5F6B7A"
LIGHT = "F3F7FB"
LIGHT_BLUE = "E8F1FA"
LIGHT_GREEN = "EAF6EF"
LIGHT_ORANGE = "FFF3E8"
LIGHT_RED = "FDEEEE"
WHITE = "FFFFFF"
BORDER = "C9D5E3"


def rgb(hex_value: str) -> tuple[int, int, int]:
    return tuple(int(hex_value[i : i + 2], 16) for i in (0, 2, 4))


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REGULAR, size)


def rounded_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    fill: str,
    outline: str = BORDER,
    radius: int = 20,
    width: int = 3,
) -> None:
    draw.rounded_rectangle(xy, radius=radius, fill=rgb(fill), outline=rgb(outline), width=width)


def center_text(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    text: str,
    size: int = 30,
    color: str = INK,
    bold: bool = False,
    spacing: int = 8,
) -> None:
    font = pil_font(size, bold)
    box = draw.multiline_textbbox((0, 0), text, font=font, spacing=spacing, align="center")
    width = box[2] - box[0]
    height = box[3] - box[1]
    x = (xy[0] + xy[2] - width) / 2
    y = (xy[1] + xy[3] - height) / 2 - box[1]
    draw.multiline_text((x, y), text, font=font, fill=rgb(color), spacing=spacing, align="center")


def arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    color: str = BLUE,
    width: int = 5,
    dashed: bool = False,
) -> None:
    if dashed:
        steps = 14
        for i in range(0, steps, 2):
            x1 = start[0] + (end[0] - start[0]) * i / steps
            y1 = start[1] + (end[1] - start[1]) * i / steps
            x2 = start[0] + (end[0] - start[0]) * (i + 1) / steps
            y2 = start[1] + (end[1] - start[1]) * (i + 1) / steps
            draw.line((x1, y1, x2, y2), fill=rgb(color), width=width)
    else:
        draw.line((*start, *end), fill=rgb(color), width=width)
    import math

    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    spread = 0.55
    p1 = (end[0] - length * math.cos(angle - spread), end[1] - length * math.sin(angle - spread))
    p2 = (end[0] - length * math.cos(angle + spread), end[1] - length * math.sin(angle + spread))
    draw.polygon([end, p1, p2], fill=rgb(color))


def title_band(draw: ImageDraw.ImageDraw, title: str, width: int) -> None:
    draw.rectangle((0, 0, width, 90), fill=rgb(NAVY))
    draw.text((45, 21), title, font=pil_font(40, True), fill=rgb(WHITE))


def save_canvas(image: Image.Image, name: str) -> Path:
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    path = IMAGE_DIR / name
    image.save(path, dpi=(180, 180), optimize=True)
    return path


def draw_architecture() -> Path:
    w, h = 1900, 1200
    im = Image.new("RGB", (w, h), rgb(WHITE))
    d = ImageDraw.Draw(im)
    title_band(d, "app-data-mcp 目标技术架构", w)

    groups = [
        ((50, 125, 330, 1060), "用户与接入", LIGHT_BLUE, BLUE),
        ((365, 125, 1160, 1060), "app-data-mcp 控制面", LIGHT, NAVY),
        ((1195, 125, 1525, 1060), "治理与元数据平面", LIGHT_GREEN, GREEN),
        ((1560, 125, 1850, 1060), "数据与分析引擎", LIGHT_ORANGE, ORANGE),
    ]
    for rect, label, fill, outline in groups:
        rounded_box(d, rect, fill, outline, 24, 4)
        d.text((rect[0] + 20, rect[1] + 15), label, font=pil_font(28, True), fill=rgb(outline))

    left_boxes = [
        ((85, 220, 295, 350), "产品 / 运营\n数据同学"),
        ((85, 420, 295, 575), "Codex / Claude\nWorkBuddy / MCP Client"),
        ((85, 660, 295, 795), "MCP HTTP / stdio"),
        ((85, 860, 295, 985), "个人 Token\n请求上下文"),
    ]
    for rect, text in left_boxes:
        rounded_box(d, rect, WHITE, BLUE)
        center_text(d, rect, text, 25, INK, True)
    for a, b in zip(left_boxes, left_boxes[1:]):
        arrow(d, ((a[0][0] + a[0][2]) // 2, a[0][3]), ((b[0][0] + b[0][2]) // 2, b[0][1]))

    control_boxes = [
        ((410, 215, 655, 345), "动态工具目录\n开关 + 用户授权"),
        ((720, 215, 960, 345), "意图识别\n结构抽取"),
        ((410, 425, 655, 565), "资产检索\n过滤 + 召回"),
        ((720, 425, 960, 565), "意图重排\n适用性评估"),
        ((410, 650, 655, 800), "资产选择治理\n预计算 Card → Metric\n→ Card → Model"),
        ((720, 650, 960, 800), "语义编译\n字段 / 参数校验"),
        ((410, 875, 655, 1000), "执行路由\n连接器"),
        ((720, 875, 960, 1000), "限行 / 限时 / 限字节\n结果标准化 + 审计"),
    ]
    for rect, text in control_boxes:
        rounded_box(d, rect, WHITE, NAVY)
        center_text(d, rect, text, 24, INK, True)
    control_edges = [(0, 1), (1, 3), (3, 2), (2, 4), (4, 5), (5, 6), (6, 7)]
    for i, j in control_edges:
        a, b = control_boxes[i][0], control_boxes[j][0]
        if abs(a[1] - b[1]) < 30:
            start, end = (a[2], (a[1] + a[3]) // 2), (b[0], (b[1] + b[3]) // 2)
        else:
            start, end = ((a[0] + a[2]) // 2, a[3]), ((b[0] + b[2]) // 2, b[1])
        arrow(d, start, end, NAVY)

    meta_boxes = [
        ((1230, 220, 1490, 350), "资产同步池\nmetadata"),
        ((1230, 420, 1490, 550), "人工语义覆盖\nadmin_overrides"),
        ((1230, 620, 1490, 750), "开放 / 有效状态\n版本与质量"),
        ((1230, 820, 1490, 980), "工具与用户授权\nSession 哈希\n审计日志"),
    ]
    for rect, text in meta_boxes:
        rounded_box(d, rect, WHITE, GREEN)
        center_text(d, rect, text, 24, INK, True)

    engine_boxes = [
        ((1595, 230, 1815, 400), "Metabase\nDashboard / Metric\nCard / Model"),
        ((1595, 500, 1815, 650), "PostHog\nDashboard / Insight"),
        ((1595, 760, 1815, 910), "StarRocks\n仅受控回退"),
    ]
    for rect, text in engine_boxes:
        rounded_box(d, rect, WHITE, ORANGE)
        center_text(d, rect, text, 24, INK, True)

    arrow(d, (295, 925), (410, 280), CYAN)
    arrow(d, (960, 495), (1230, 285), GREEN)
    arrow(d, (960, 720), (1230, 685), GREEN)
    arrow(d, (960, 940), (1230, 900), GREEN)
    arrow(d, (960, 940), (1595, 315), ORANGE)
    arrow(d, (960, 940), (1595, 575), ORANGE)
    arrow(d, (960, 940), (1595, 835), RED, dashed=True)
    d.text((1565, 1030), "管理后台维护开放、语义、工具与审计", font=pil_font(22, True), fill=rgb(MUTED))
    return save_canvas(im, "01-target-architecture.png")


def draw_asset_pool() -> Path:
    w, h = 1800, 720
    im = Image.new("RGB", (w, h), rgb(WHITE))
    d = ImageDraw.Draw(im)
    title_band(d, "资产从平台同步到 AI 开放的治理漏斗", w)
    boxes = [
        ((55, 260, 280, 455), "平台全量对象\nMetabase / PostHog", LIGHT_BLUE, BLUE),
        ((355, 260, 590, 455), "只读同步池\n全量元信息\n新资产默认关闭", LIGHT, NAVY),
        ((675, 245, 930, 470), "准入检查\nOwner / 口径 / 字段\n权限 / 质量 / 可运行", LIGHT_ORANGE, ORANGE),
        ((1035, 180, 1300, 375), "精选开放集\n20+ 指标\n30+ 数据集", LIGHT_GREEN, GREEN),
        ((1035, 465, 1300, 635), "待治理 / 关闭\n重复、过期或信息不全", LIGHT_RED, RED),
        ((1405, 180, 1735, 375), "AI 可检索与执行\n持续记录使用与质量", LIGHT_BLUE, BLUE),
    ]
    for rect, text, fill, outline in boxes:
        rounded_box(d, rect, fill, outline, 24, 4)
        center_text(d, rect, text, 27, INK, True)
    arrow(d, (280, 357), (355, 357))
    arrow(d, (590, 357), (675, 357))
    arrow(d, (930, 330), (1035, 280), GREEN)
    arrow(d, (930, 390), (1035, 545), RED)
    arrow(d, (1300, 280), (1405, 280))
    arrow(d, (1570, 375), (1180, 465), MUTED, dashed=True)
    d.text((1340, 435), "异常反馈 / 复核 / 退役", font=pil_font(22, True), fill=rgb(MUTED))
    return save_canvas(im, "02-asset-governance-funnel.png")


def draw_request_flow() -> Path:
    w, h = 1700, 1800
    im = Image.new("RGB", (w, h), rgb(WHITE))
    d = ImageDraw.Draw(im)
    title_band(d, "自然语言问题到受控数据结果的详细处理流程", w)
    steps = [
        (120, "1  用户自然语言问题", LIGHT_BLUE, BLUE),
        (285, "2  search_assets：原始问题必须先检索", LIGHT, NAVY),
        (450, "3  Token 校验 + 发布状态 + 权限快照过滤", LIGHT_GREEN, GREEN),
        (615, "4  意图识别与结构抽取\n指标 / 时间 / 维度 / 筛选 / 实体 / 动作", LIGHT, NAVY),
        (800, "5  混合召回 + 意图重排 + 适用性解释", LIGHT_BLUE, BLUE),
        (965, "6  get_asset 核验公式、字段、粒度、参数与警告", LIGHT_ORANGE, ORANGE),
        (1130, "7  资产满足问题？\n是：run_asset　否：拒绝候选并检查下一项", LIGHT, NAVY),
        (1315, "8  服务端二次治理 + 实时权限 + 语义/参数校验", LIGHT_GREEN, GREEN),
        (1480, "9  只读执行 + 超时/行数/字节保护", LIGHT_RED, RED),
        (1645, "10  标准化结果 + 来源 + 警告 + 审计 → AI 解释", LIGHT_BLUE, BLUE),
    ]
    rects = []
    for y, text, fill, outline in steps:
        rect = (280, y, 1420, y + 115)
        rects.append(rect)
        rounded_box(d, rect, fill, outline, 24, 4)
        center_text(d, rect, text, 27, INK, True)
    for a, b in zip(rects, rects[1:]):
        arrow(d, ((a[0] + a[2]) // 2, a[3]), ((b[0] + b[2]) // 2, b[1]), NAVY)

    side = (60, 1060, 235, 1350)
    rounded_box(d, side, LIGHT_RED, RED, 22, 4)
    center_text(d, side, "治理资产\n确实不适用\n且策略允许\n↓\n受控 SQL 回退", 23, RED, True)
    arrow(d, (280, 1185), (235, 1185), RED)
    arrow(d, (150, 1350), (280, 1540), RED, dashed=True)
    note = (1460, 1030, 1640, 1370)
    rounded_box(d, note, LIGHT_ORANGE, ORANGE, 22, 4)
    center_text(d, note, "硬规则\n\nMetric / Card\n优先\n\n不得让 AI\n静默换口径", 23, ORANGE, True)
    return save_canvas(im, "03-end-to-end-request-flow.png")


def draw_semantic_lifecycle() -> Path:
    w, h = 1800, 760
    im = Image.new("RGB", (w, h), rgb(WHITE))
    d = ImageDraw.Draw(im)
    title_band(d, "语义资产发布与持续治理生命周期", w)
    nodes = [
        ((80, 285, 290, 445), "Synced\n平台同步", LIGHT_BLUE, BLUE),
        ((370, 285, 580, 445), "Draft\n补齐语义", LIGHT, NAVY),
        ((660, 285, 895, 445), "Reviewing\nSchema / 样例\nOwner 评审", LIGHT_ORANGE, ORANGE),
        ((1000, 285, 1230, 445), "Published\n进入精选集", LIGHT_GREEN, GREEN),
        ((1335, 170, 1575, 330), "Suspended\n质量 / 权限异常", LIGHT_RED, RED),
        ((1335, 465, 1575, 625), "Retired\n废弃 / 平台删除", LIGHT, MUTED),
    ]
    for rect, text, fill, outline in nodes:
        rounded_box(d, rect, fill, outline, 24, 4)
        center_text(d, rect, text, 26, INK, True)
    arrow(d, (290, 365), (370, 365))
    arrow(d, (580, 365), (660, 365))
    arrow(d, (895, 365), (1000, 365), GREEN)
    arrow(d, (780, 445), (475, 445), ORANGE)
    d.text((555, 470), "评审退回", font=pil_font(21, True), fill=rgb(ORANGE))
    arrow(d, (1230, 320), (1335, 250), RED)
    arrow(d, (1455, 330), (1170, 445), GREEN, dashed=True)
    d.text((1250, 355), "修复并回归", font=pil_font(21, True), fill=rgb(GREEN))
    arrow(d, (1230, 410), (1335, 545), MUTED)
    d.text((1060, 535), "删除 / 口径废弃", font=pil_font(21, True), fill=rgb(MUTED))
    d.text((100, 650), "发布门槛：Owner + Reviewer + 结构化语义 + 3 个正例问题 + 1 个反例 + 查询回归通过", font=pil_font(26, True), fill=rgb(NAVY))
    return save_canvas(im, "04-semantic-lifecycle.png")


def draw_permission_sequence() -> Path:
    w, h = 1900, 1050
    im = Image.new("RGB", (w, h), rgb(WHITE))
    d = ImageDraw.Draw(im)
    title_band(d, "个人 Token 到 BI 最终权限裁决的三道权限门", w)
    actors = [("用户", 160), ("AI 客户端", 520), ("app-data-mcp", 930), ("元数据 PostgreSQL", 1360), ("Metabase", 1740)]
    for label, x in actors:
        rounded_box(d, (x - 120, 130, x + 120, 220), LIGHT_BLUE if x < 900 else LIGHT_GREEN, BLUE if x < 900 else GREEN)
        center_text(d, (x - 120, 130, x + 120, 220), label, 23, INK, True)
        d.line((x, 220, x, 960), fill=rgb(BORDER), width=3)
    messages = [
        (270, 160, 520, "配置个人 MCP token", BLUE),
        (365, 520, 930, "Authorization: Bearer appdata_xxx", BLUE),
        (460, 930, 930, "Token 哈希映射用户与 Session", NAVY),
        (555, 930, 1360, "门 1+2：仅查 published / active\n并按权限快照过滤", GREEN),
        (675, 930, 1740, "门 3：个人 Session 实时校验资产", ORANGE),
        (780, 1740, 930, "允许 / 401 / 403 / 404", ORANGE),
        (875, 930, 1740, "同一 Session 执行只读查询", GREEN),
        (970, 1740, 520, "权限范围内结果 + 来源 + 警告", BLUE),
    ]
    for y, x1, x2, text, color in messages:
        arrow(d, (x1, y), (x2, y), color, 4, dashed=(x1 == x2))
        tx = min(x1, x2) + abs(x2 - x1) // 2
        bbox = d.multiline_textbbox((0, 0), text, font=pil_font(20, True), spacing=5, align="center")
        d.multiline_text((tx - (bbox[2] - bbox[0]) / 2, y - 48), text, font=pil_font(20, True), fill=rgb(color), spacing=5, align="center")
    return save_canvas(im, "05-permission-sequence.png")


def draw_kpi_chart() -> Path:
    w, h = 1700, 900
    im = Image.new("RGB", (w, h), rgb(WHITE))
    d = ImageDraw.Draw(im)
    title_band(d, "Q3 核心交付目标（目标值，不代表当前完成度）", w)
    labels = ["治理指标", "治理数据集", "非数据团队用户", "周活跃查询"]
    values = [20, 30, 20, 50]
    colors = [BLUE, GREEN, CYAN, ORANGE]
    base_y, max_h = 720, 500
    chart_left, chart_right = 160, 1580
    d.line((chart_left, base_y, chart_right, base_y), fill=rgb(INK), width=4)
    for tick in range(0, 51, 10):
        y = base_y - tick / 50 * max_h
        d.line((chart_left, y, chart_right, y), fill=rgb(BORDER), width=2)
        d.text((90, y - 15), str(tick), font=pil_font(20), fill=rgb(MUTED))
    bar_w, gap = 210, 125
    x = 250
    for label, value, color in zip(labels, values, colors):
        height = value / 50 * max_h
        d.rounded_rectangle((x, base_y - height, x + bar_w, base_y), radius=18, fill=rgb(color))
        val_text = f"≥ {value}" if value != 50 else "≥ 50/周"
        box = d.textbbox((0, 0), val_text, font=pil_font(34, True))
        d.text((x + (bar_w - box[2]) / 2, base_y - height - 58), val_text, font=pil_font(34, True), fill=rgb(color))
        box2 = d.textbbox((0, 0), label, font=pil_font(24, True))
        d.text((x + (bar_w - box2[2]) / 2, base_y + 25), label, font=pil_font(24, True), fill=rgb(INK))
        x += bar_w + gap
    d.text((160, 820), "另：常见查询从提问到结果 ≤ 1 分钟；9 月底完成中心现场验收。", font=pil_font(26, True), fill=rgb(NAVY))
    return save_canvas(im, "06-q3-kpi-targets.png")


def draw_roadmap() -> Path:
    w, h = 1900, 1120
    im = Image.new("RGB", (w, h), rgb(WHITE))
    d = ImageDraw.Draw(im)
    title_band(d, "app-data-mcp Q3 实施 Roadmap（2026）", w)
    left, right = 510, 1830
    start_day, total_days = 0, 66  # 7/27—9/30
    milestones = [
        ("基线与规范", "资产准入、语义 Schema、黄金问题集", 0, 14, NAVY),
        ("平台与语义", "Metabase 语义治理增强", 7, 25, BLUE),
        ("平台与语义", "PostHog Insight 语义与表单", 14, 32, CYAN),
        ("检索", "混合召回、覆盖评分、可解释重排", 14, 34, PURPLE if "PURPLE" in globals() else NAVY),
        ("安全可靠性", "Token、敏感字段、告警与故障演练", 21, 43, ORANGE),
        ("资产治理", "首批 10 指标 + 15 数据集", 7, 26, GREEN),
        ("资产治理", "达成 20+ 指标 + 30+ 数据集", 28, 47, GREEN),
        ("试点", "5 人 Alpha → 20+ 人 Beta / 周活 50+", 28, 58, CYAN),
        ("验收", "准确性、权限、稳定性回归与压测", 49, 60, RED),
        ("验收", "中心现场验收与整改", 60, 66, NAVY),
    ]
    d.text((40, 135), "工作流 / 交付", font=pil_font(25, True), fill=rgb(NAVY))
    ticks = [(0, "7/27"), (10, "8/06"), (20, "8/16"), (30, "8/26"), (40, "9/05"), (50, "9/15"), (60, "9/25"), (66, "9/30")]
    for day, label in ticks:
        x = left + day / total_days * (right - left)
        d.line((x, 175, x, 1040), fill=rgb(BORDER), width=2)
        d.text((x - 28, 135), label, font=pil_font(20, True), fill=rgb(MUTED))
    y = 215
    for group, label, start, end, color in milestones:
        d.text((40, y + 8), group, font=pil_font(21, True), fill=rgb(color))
        d.text((170, y + 8), label, font=pil_font(21), fill=rgb(INK))
        x1 = left + start / total_days * (right - left)
        x2 = left + end / total_days * (right - left)
        d.rounded_rectangle((x1, y, x2, y + 48), radius=14, fill=rgb(color))
        y += 78
    milestone_x = left + 60 / total_days * (right - left)
    d.polygon([(milestone_x, 985), (milestone_x + 18, 1003), (milestone_x, 1021), (milestone_x - 18, 1003)], fill=rgb(RED))
    d.text((milestone_x - 70, 1030), "9/25 现场验收", font=pil_font(20, True), fill=rgb(RED))
    return save_canvas(im, "07-q3-roadmap.png")


def set_run_font(run, size: float = 11, bold: bool = False, color: str = INK, name: str = "STHeiti", italic: bool = False) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    rpr.rFonts.set(qn("w:eastAsia"), name)
    rpr.rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=BORDER, size=6) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for name, value in (("tblW", sum(widths_dxa)), ("tblInd", indent_dxa)):
        node = tbl_pr.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tbl_pr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            set_cell_borders(cell)


def add_numbering(doc: Document) -> None:
    numbering = doc.part.numbering_part.element
    for abstract_id, num_id, fmt, text_value in ((100, 100, "bullet", "•"), (101, 101, "decimal", "%1.")):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text_value)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        lvl.extend([start, num_fmt, lvl_text, p_pr])
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(abstract_id))
        num.append(abs_id)
        numbering.append(num)
    doc._bullet_num_id = 100
    doc._decimal_num_id = 101
    doc._current_decimal_num_id = 101
    doc._next_num_id = 102


def restart_decimal_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    num_id = doc._next_num_id
    doc._next_num_id += 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), "101")
    num.append(abs_id)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    doc._current_decimal_num_id = num_id
    return num_id


def set_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "STHeiti"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "STHeiti")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "STHeiti")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, NAVY, 8, 4),
    }
    for style_name, (size, color, before, after) in tokens.items():
        style = doc.styles[style_name]
        style.font.name = "STHeiti"
        style._element.rPr.rFonts.set(qn("w:ascii"), "STHeiti")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "STHeiti")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "STHeiti"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_numbering(doc)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(2)
    set_run_font(hp.add_run("APP-DATA-MCP · Q3 技术实现方案"), 9, True, MUTED)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(fp.add_run("内部讨论稿  |  "), 9, False, MUTED)
    add_field(fp, "PAGE")


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run("TECHNICAL IMPLEMENTATION PLAN"), 10, True, CYAN)
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    set_run_font(title.add_run("app-data-mcp\n技术实现方案"), 28, True, NAVY)
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(22)
    set_run_font(sub.add_run("面向 Q3 数据自助分析目标的架构、语义治理、安全与实施路线"), 14, False, BLUE)

    for label, value in (
        ("版本", "V1.0（讨论稿）"),
        ("日期", "2026-07-24"),
        ("范围", "Metabase · PostHog · MCP · 个人权限 · 管理后台"),
        ("受众", "数据 / 后端 / 产品 / 运营 / 信息安全 / 管理层"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        set_run_font(p.add_run(f"{label}："), 10.5, True, MUTED)
        set_run_font(p.add_run(value), 10.5, False, INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.14)
    p.paragraph_format.line_spacing = 1.2
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_BLUE)
    p_pr.append(shd)
    set_run_font(p.add_run("方案定位  "), 11, True, BLUE)
    set_run_font(
        p.add_run("位于 AI 工具与 BI/数据引擎之间的只读数据控制面：用精选资产、结构化语义、个人权限、执行护栏和审计闭环，把自然语言问题稳定转换成可解释、可追溯、可控制的数据查询。"),
        11,
        False,
        INK,
    )

    nav = doc.add_paragraph()
    nav.paragraph_format.space_before = Pt(18)
    nav.paragraph_format.space_after = Pt(6)
    set_run_font(nav.add_run("阅读导航"), 13, True, NAVY)
    for text in (
        "管理层：先读第 1、3、9、10、15 节",
        "数据与产品：重点读第 4、5、6、9、12 节",
        "后端与安全：重点读第 3、5、7、8、10、13 节",
    ):
        p = doc.add_paragraph()
        set_num(p, doc._bullet_num_id)
        set_run_font(p.add_run(text), 10.5)
    doc.add_page_break()


def add_inline_markdown(paragraph, text: str, default_size: float = 11, default_color: str = INK) -> None:
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            set_run_font(paragraph.add_run(text[pos : match.start()]), default_size, False, default_color)
        token = match.group(0)
        if token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), default_size, True, default_color)
        else:
            set_run_font(paragraph.add_run(token[1:-1]), default_size - 0.3, False, NAVY, "STHeiti")
        pos = match.end()
    if pos < len(text):
        set_run_font(paragraph.add_run(text[pos:]), default_size, False, default_color)


def add_callout(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.18
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_BLUE)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BLUE)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    add_inline_markdown(p, text, 10.5)


def add_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.16)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_together = True
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "182332")
    p_pr.append(shd)
    lines = code.splitlines()
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, 8.4, False, "EAF2FA", "STHeiti")
        if idx < len(lines) - 1:
            run.add_break()


def table_widths(rows: list[list[str]]) -> list[int]:
    cols = len(rows[0])
    if cols == 2:
        return [2500, 6860]
    if cols == 3:
        return [2150, 2150, 5060]
    if cols == 4:
        return [1850, 1450, 2450, 3610]
    if cols == 5:
        return [1300, 1500, 1800, 1800, 2960]
    base = 9360 // cols
    widths = [base] * cols
    widths[-1] += 9360 - sum(widths)
    return widths


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    widths = table_widths(rows)
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if i == 0:
                shade_cell(cell, LIGHT)
                set_run_font(p.add_run(value), 9.2, True, NAVY)
            else:
                set_run_font(p.add_run(value), 9.1, False, INK)
            if j > 0 and len(value) < 12:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_geometry(table, widths, 120)
    header_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_pr.append(repeat)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc: Document, path: Path, caption: str, alt: str, width: float = 6.35) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", alt)
    shape._inline.docPr.set("title", caption)
    cp = doc.add_paragraph(style="Caption")
    set_run_font(cp.add_run(caption), 9, False, MUTED)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def render_markdown(doc: Document, image_paths: list[Path]) -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    diagram_captions = [
        ("图 1　app-data-mcp 目标技术架构", "AI 接入、控制面、元数据治理平面与数据引擎之间的关系"),
        ("图 2　同步池与精选开放集的资产治理漏斗", "从全量同步到准入、开放、反馈和退役"),
        ("图 3　自然语言问题到数据结果的端到端流程", "包含检索、核验、治理、权限、执行与审计"),
        ("图 4　语义资产发布生命周期", "同步、草稿、评审、发布、暂停和退役状态"),
        ("图 5　个人权限校验时序", "个人 Token、开放状态、权限快照与 Metabase 实时裁决"),
        ("图 6　Q3 核心交付目标", "20+ 指标、30+ 数据集、20+ 用户与每周 50+ 查询"),
        ("图 7　Q3 实施 Roadmap", "从规范、语义和资产建设到试点与现场验收"),
    ]
    diagram_index = 0
    i = 0
    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = stripped[3:].strip()
                code_lines = []
            else:
                if code_lang in {"mermaid"}:
                    caption, alt = diagram_captions[diagram_index]
                    add_figure(doc, image_paths[diagram_index], caption, alt)
                    diagram_index += 1
                else:
                    add_code(doc, "\n".join(code_lines))
                in_code = False
                code_lang = ""
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith("# "):
            i += 1
            continue
        if stripped.startswith("## "):
            if stripped.startswith("## 15."):
                doc.add_page_break()
            p = doc.add_paragraph(style="Heading 1")
            add_inline_markdown(p, stripped[3:], 16, BLUE)
            i += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline_markdown(p, stripped[4:], 13, BLUE)
            i += 1
            continue
        if stripped.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline_markdown(p, stripped[5:], 12, NAVY)
            i += 1
            continue
        if stripped.startswith("> "):
            add_callout(doc, stripped[2:])
            i += 1
            continue
        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        if re.match(r"^\d+\.\s+", stripped):
            source_number = int(re.match(r"^(\d+)\.", stripped).group(1))
            if source_number == 1:
                restart_decimal_numbering(doc)
            p = doc.add_paragraph()
            set_num(p, doc._current_decimal_num_id)
            add_inline_markdown(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph()
            set_num(p, doc._bullet_num_id)
            add_inline_markdown(p, stripped[2:])
            i += 1
            continue
        if stripped.startswith("**") and stripped.endswith("**"):
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            add_inline_markdown(p, stripped, 11, NAVY)
            i += 1
            continue
        p = doc.add_paragraph()
        add_inline_markdown(p, stripped)
        i += 1


def set_core_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "app-data-mcp 技术实现方案（Q3）"
    props.subject = "自然语言数据自助分析的架构、语义治理、安全与实施路线"
    props.author = "app-data-mcp 项目组"
    props.keywords = "MCP, Metabase, PostHog, 数据语义, 权限, Q3 Roadmap"
    props.comments = "基于 2026-07-24 代码仓库现状形成的讨论稿。"


def build() -> None:
    image_paths = [
        draw_architecture(),
        draw_asset_pool(),
        draw_request_flow(),
        draw_semantic_lifecycle(),
        draw_permission_sequence(),
        draw_kpi_chart(),
        draw_roadmap(),
    ]
    doc = Document()
    configure_document(doc)
    add_header_footer(doc)
    set_core_properties(doc)
    add_cover(doc)
    render_markdown(doc, image_paths)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
