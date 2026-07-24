from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/lute/code/app-data-mcp")
OUT = ROOT / "docs" / "app-data-mcp架构介绍、配置接入及使用指引-微软雅黑版.docx"
ARCH = ROOT / "docs" / "images" / "app-data-mcp-architecture.png"
IMG_AUTH = ROOT / "docs" / "images" / "app-data-mcp-auth-page.png"
IMG_QUERY = ROOT / "docs" / "images" / "workbuddy-single-metric-us-bind-users.png"
IMG_AUDIENCE_ANALYSIS = ROOT / "docs" / "images" / "workbuddy-audience-intersection-analysis.png"
IMG_DEVICE_GROUP_ANALYSIS = ROOT / "docs" / "images" / "workbuddy-device-model-distribution.png"
IMG_APP_BIND_TREND_ANALYSIS = ROOT / "docs" / "images" / "workbuddy-app-bind-penetration-trend.png"
IMG_APP_BIND_DUAL_AXIS_TREND = ROOT / "docs" / "images" / "workbuddy-app-bind-dual-axis-trend.png"
IMG_POSTHOG_INSIGHT_INVENTORY = ROOT / "docs" / "images" / "workbuddy-posthog-insight-inventory.png"
IMG_POSTHOG_ANALYSIS_REPORT = ROOT / "docs" / "images" / "workbuddy-posthog-analysis-report.png"

IMG_WB_CONNECTOR = ROOT / "docs" / "images" / "workbuddy-step-1-open-connectors.png"
IMG_WB_CONFIG = ROOT / "docs" / "images" / "workbuddy-step-2-configure-mcp.png"
IMG_WB_ENABLE = ROOT / "docs" / "images" / "workbuddy-step-3-enable-app-data.png"

NAVY = "123B6D"
BLUE = "1D67B1"
CYAN = "0A91B8"
GREEN = "17864B"
ORANGE = "D66A1F"
PURPLE = "6B43C3"
INK = "172033"
MUTED = "5F6B7A"
LIGHT = "F3F7FB"
LIGHT_BLUE = "E8F1FA"
LIGHT_GREEN = "EAF6EF"
LIGHT_ORANGE = "FFF3E8"
LIGHT_RED = "FDEEEE"
WHITE = "FFFFFF"
BORDER = "C9D5E3"


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=BORDER, size=6):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    trPr.append(cant_split)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths_dxa)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(indent_dxa))
    tblInd.set(qn("w:type"), "dxa")
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[idx]))
            tcW.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, size=11, bold=False, color=INK, name="Microsoft YaHei", italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = instruction
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    rFonts.set(qn("w:cs"), "Microsoft YaHei")
    rPr.extend([rFonts, color, underline])
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_paragraph_shading(paragraph, fill, border_color=None):
    pPr = paragraph._p.get_or_add_pPr()
    shd = pPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        pPr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border_color:
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is None:
            pBdr = OxmlElement("w:pBdr")
            pPr.append(pBdr)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border_color)
        pBdr.append(left)


def add_alt_text(inline_shape, descr):
    docPr = inline_shape._inline.docPr
    docPr.set("descr", descr)
    docPr.set("title", descr)


def add_figure(doc, image_path, caption, width=6.2, alt=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(image_path), width=Inches(width))
    add_alt_text(shape, alt or caption)
    cp = doc.add_paragraph(style="Caption")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.keep_with_next = False
    set_run_font(cp.add_run(caption), size=9, color=MUTED)
    return shape


def add_callout(doc, label, text, kind="info"):
    colors = {
        "info": (LIGHT_BLUE, BLUE),
        "success": (LIGHT_GREEN, GREEN),
        "warn": (LIGHT_ORANGE, ORANGE),
        "risk": (LIGHT_RED, "B42318"),
    }
    fill, accent = colors[kind]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.18
    add_paragraph_shading(p, fill, accent)
    set_run_font(p.add_run(f"{label}  "), size=10.5, bold=True, color=accent)
    set_run_font(p.add_run(text), size=10.5, color=INK)
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    add_paragraph_shading(p, "182332", "3D526A")
    lines = code.splitlines()
    for i, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, size=9.2, color="EAF2FA", name="Menlo")
        if i < len(lines) - 1:
            run.add_break()
    return p


def add_prompt(doc, title, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.keep_together = True
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.10)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.18
    add_paragraph_shading(p, "F4F7FB", CYAN)
    set_run_font(p.add_run(f"{title}\n"), size=10.5, bold=True, color=NAVY)
    set_run_font(p.add_run(text), size=10.2, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(text), size={1: 16, 2: 13, 3: 12}[level], bold=True,
                 color={1: BLUE, 2: NAVY, 3: CYAN}[level])
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), bold=True, color=INK)
        set_run_font(p.add_run(text[len(bold_prefix):]), color=INK)
    else:
        set_run_font(p.add_run(text), color=INK)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="MCP Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.63)
        p.paragraph_format.first_line_indent = Inches(-0.18)
    set_run_font(p.add_run(text), color=INK)
    return p


def add_number(doc, text, restart=False):
    if restart:
        numbering = doc.part.numbering_part.element
        next_id = max([int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))] or [0]) + 1
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(next_id))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(doc._mcp_number_abstract_id))
        num.append(abs_id)
        lvl_override = OxmlElement("w:lvlOverride")
        lvl_override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        lvl_override.append(start_override)
        num.append(lvl_override)
        numbering.append(num)
        doc._mcp_current_number_id = next_id
    p = doc.add_paragraph(style="MCP Number")
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), str(doc._mcp_current_number_id))
    numPr.extend([ilvl, numId])
    p._p.get_or_add_pPr().append(numPr)
    set_run_font(p.add_run(text), color=INK)
    return p


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for idx, before, after, size, color in (
        (1, 18, 10, 16, BLUE), (2, 14, 7, 13, NAVY), (3, 10, 5, 12, CYAN)
    ):
        style = styles[f"Heading {idx}"]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:cs"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Microsoft YaHei"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption._element.rPr.rFonts.set(qn("w:cs"), "Microsoft YaHei")
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_after = Pt(7)

    numbering = doc.part.numbering_part.element
    max_abs = max([int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))] or [0])
    max_num = max([int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))] or [0])

    def create_num(num_format, text, aligned, indent, hanging):
        nonlocal max_abs, max_num
        max_abs += 1
        max_num += 1
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(max_abs))
        nsid = OxmlElement("w:nsid")
        nsid.set(qn("w:val"), f"A1B2{max_abs:04X}"[-8:])
        abstract.append(nsid)
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        numFmt = OxmlElement("w:numFmt")
        numFmt.set(qn("w:val"), num_format)
        lvlText = OxmlElement("w:lvlText")
        lvlText.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        pPr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(indent))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(indent))
        ind.set(qn("w:hanging"), str(hanging))
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        pPr.extend([tabs, ind, spacing])
        lvl.extend([start, numFmt, lvlText, suff, pPr])
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(max_num))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(max_abs))
        num.append(abs_id)
        numbering.append(num)
        return max_num, max_abs

    bullet_id, _bullet_abs = create_num("bullet", "•", 260, 540, 270)
    number_id, number_abs = create_num("decimal", "%1.", 260, 540, 270)
    doc._mcp_number_abstract_id = number_abs
    doc._mcp_current_number_id = number_id
    for name, num_id in (("MCP Bullet", bullet_id), ("MCP Number", number_id)):
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        numPr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        numId = OxmlElement("w:numId")
        numId.set(qn("w:val"), str(num_id))
        numPr.extend([ilvl, numId])
        style.element.get_or_add_pPr().append(numPr)


def add_header_footer(section, first=False):
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = first
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("APP生态中心 · app-data-mcp 使用指南"), size=8.5, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    set_run_font(fp.add_run("内部使用  |  "), size=8.5, color=MUTED)
    add_field(fp, "PAGE")


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build():
    for path in (ARCH, IMG_AUTH, IMG_QUERY, IMG_AUDIENCE_ANALYSIS,
                 IMG_DEVICE_GROUP_ANALYSIS, IMG_APP_BIND_TREND_ANALYSIS,
                 IMG_POSTHOG_INSIGHT_INVENTORY, IMG_POSTHOG_ANALYSIS_REPORT,

                 IMG_APP_BIND_DUAL_AXIS_TREND,
                 IMG_WB_CONNECTOR, IMG_WB_CONFIG, IMG_WB_ENABLE):
        if not path.exists():
            raise FileNotFoundError(path)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    configure_styles(doc)
    add_header_footer(section, first=True)

    # Cover: compact_reference_guide + editorial_cover opening pattern.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(58)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("APP生态中心 · 数据平台"), size=11, bold=True, color=CYAN)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    set_run_font(p.add_run("app-data-mcp架构介绍、\n配置接入及使用指引"), size=27, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    set_run_font(p.add_run("让每位同事都能通过 AI 工具安全、准确地获取和分析日常数据"), size=13, color=MUTED)
    add_callout(doc, "文档目标", "说明 app-data-mcp 为什么开放、如何接入、怎样正确提问，以及如何从一次查询走向可复用的深入分析。", "info")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("版本 1.0  |  2026年7月  |  内部使用"), size=10, color=MUTED)
    page_break(doc)

    add_heading(doc, "内容导航", 1)
    nav = [
        "1. 背景与开放目标",
        "2. app-data-mcp 架构介绍",
        "3. WorkBuddy 配置接入",
        "4. 使用指引：从连接验证到深入分析",
        "5. 安全规范、常见问题与快速检查清单",
    ]
    for item in nav:
        add_bullet(doc, item)
    add_callout(doc, "建议阅读方式", "首次接入的同事依次阅读第 3～5 章；已经完成配置的同事可先使用“首次连接验证提示词”确认接入，再从“高质量提问模板”开始正式分析。", "success")

    add_heading(doc, "1. 背景与开放目标", 1)
    add_heading(doc, "1.1 为什么开放 app-data-mcp", 2)
    add_body(doc, "APP生态中心的产品、运营和业务团队每天都会遇到用户规模、设备绑定、活跃趋势、行为漏斗、人群交集等分析需求。以往不少问题需要由数据同学完成资产定位、口径确认、取数和解释，链路较长，也难以覆盖全部日常分析。")
    add_body(doc, "app-data-mcp 将数据开发平台（Metabase）与行为分析平台（PostHog）的数据资产和查询能力，通过统一的 MCP 服务开放给 WorkBuddy、Codex、Claude Code 等 AI 工具。业务同事可直接用自然语言查找资产、获取数据、核对口径并继续追问分析。")
    add_heading(doc, "1.2 开放后的协作方式", 2)
    add_bullet(doc, "日常分析自助化：常规指标、趋势、分组、筛选和人群查询由业务同事通过 AI 工具完成。")
    add_bullet(doc, "数据团队聚焦高价值工作：重点项目深度经营分析、复杂建模、指标治理和数据产品评审。")
    add_bullet(doc, "统一数据入口：通过 MCP 复用已治理的 Model、Metric、Card、Dashboard、Insight 等资产。")
    add_bullet(doc, "权限随用户身份：授权后按个人数据权限提供只读查询能力，不绕过现有平台权限。")
    add_callout(doc, "定位", "app-data-mcp 是“资产发现 + 权限校验 + 数据查询 + 分析辅助”的统一入口，不替代指标治理、数据质量管理或正式经营结论的评审流程。", "warn")
    page_break(doc)

    add_heading(doc, "2. app-data-mcp 架构介绍", 1)
    add_body(doc, "端到端流程可以概括为：同事提出自然语言问题，AI 工具通过 MCP 协议调用 app-data-mcp；服务完成身份认证、权限校验、资产搜索和查询编排，再访问 Metabase 或 PostHog；结果返回 AI 工具后，由同事继续进行对比、拆解、趋势和归因分析。")
    add_figure(doc, ARCH, "图 1  app-data-mcp 端到端服务架构", width=6.45,
               alt="app-data-mcp 端到端架构图，包含 AI 使用入口、MCP 服务、Metabase、PostHog 和分析闭环。")
    add_callout(doc, "主流程", "提出问题 → 搜索资产 → 选择并查询 → 返回数据 → 深入分析 → 继续追问。", "info")

    add_heading(doc, "2.1 服务分层", 2)
    rows = [
        ("使用入口", "WorkBuddy、Codex、Claude Code 等支持 MCP 的 AI 工具", "自然语言对话与多轮分析"),
        ("统一 MCP 接入层", "工具发现、参数校验、响应封装", "为不同 AI 工具提供一致调用体验"),
        ("身份认证与权限控制", "用户身份、资产权限、查询鉴权", "确保查询权限与个人账号一致"),
        ("数据资产搜索", "元信息检索、资产发现、上下文补全", "先找对资产，再执行数据查询"),
        ("查询与工具编排", "路由选择、参数转换、结果标准化", "将请求路由到 Metabase 或 PostHog"),
        ("分析结果返回", "结构化数据、口径说明、解释与洞察", "支持继续追问和深度拆解"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [1900, 3500, 3960])
    headers = ["层级", "主要能力", "对使用者的价值"]
    for i, text in enumerate(headers):
        shade_cell(table.rows[0].cells[i], NAVY)
        set_cell_borders(table.rows[0].cells[i], NAVY)
        p = table.rows[0].cells[i].paragraphs[0]
        set_run_font(p.add_run(text), size=9.5, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, text in enumerate(row):
            if idx % 2:
                shade_cell(cells[i], LIGHT)
            set_cell_borders(cells[i])
            p = cells[i].paragraphs[0]
            set_run_font(p.add_run(text), size=9.2, bold=(i == 0), color=INK)
    for row in table.rows:
        set_row_cant_split(row)
    add_heading(doc, "2.2 后端平台与数据资产", 2)
    add_body(doc, "app-data-mcp 不要求同事记住每个资产 ID。更推荐的方式是先描述业务问题和口径，让 AI 搜索候选资产、展示名称与用途，再确认查询对象。")
    assets = [
        ("数据开发平台（Metabase）", "Model", "可查询的数据模型，适合筛选、分组、聚合和明细探索"),
        ("数据开发平台（Metabase）", "Metric", "已沉淀的语义指标与统一统计口径"),
        ("数据开发平台（Metabase）", "Card", "已配置的问题、报表或固定查询"),
        ("数据开发平台（Metabase）", "Dashboard", "多个指标或报表组成的数据看板"),
        ("行为分析平台（PostHog）", "Insight", "趋势、漏斗、留存等用户行为洞察"),
        ("行为分析平台（PostHog）", "Dashboard", "行为分析看板及其关联洞察"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [2900, 1600, 4860])
    for i, text in enumerate(("平台", "资产类型", "典型用途")):
        shade_cell(table.rows[0].cells[i], BLUE)
        set_cell_borders(table.rows[0].cells[i], BLUE)
        set_run_font(table.rows[0].cells[i].paragraphs[0].add_run(text), size=9.5, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for idx, row in enumerate(assets):
        cells = table.add_row().cells
        for i, text in enumerate(row):
            shade_cell(cells[i], LIGHT if idx % 2 else WHITE)
            set_cell_borders(cells[i])
            set_run_font(cells[i].paragraphs[0].add_run(text), size=9.2, bold=(i == 1), color=INK)
    for row in table.rows:
        set_row_cant_split(row)

    add_heading(doc, "2.3 一次查询在系统中的流转", 2)
    for idx, text in enumerate((
        "同事在 AI 工具中描述问题，并明确要求使用 app-data MCP。",
        "AI 工具读取 MCP 工具说明，先搜索 Metabase/PostHog 中的相关资产。",
        "app-data-mcp 校验个人 token 和资产权限，并将请求路由到对应平台。",
        "平台返回元信息或数据；app-data-mcp 将结果标准化后交给 AI 工具。",
        "AI 展示指标、数据源、筛选条件、统计口径和数据时效，并根据追问继续分析。",
    )):
        add_number(doc, text, restart=(idx == 0))
    add_callout(doc, "重要原则", "当业务问题存在多个可能口径时，应先让 AI 展示候选资产及差异，再由使用者确认；不要让 AI 猜测字段、指标定义或过滤规则。", "warn")

    page_break(doc)
    add_heading(doc, "3. WorkBuddy 配置接入", 1)
    add_heading(doc, "3.1 接入前准备", 2)
    add_bullet(doc, "已安装可使用的 WorkBuddy，并能够修改其 MCP 配置文件。")
    add_bullet(doc, "拥有数据开发平台（Metabase）账号及对应数据权限。")
    add_bullet(doc, "能够访问授权页面和 MCP 服务地址。")
    add_bullet(doc, "确认仅在本人设备和本人账号下保存个人 MCP token。")
    add_heading(doc, "3.2 第一步：完成身份授权并获取个人 token", 2)
    add_number(doc, "在浏览器打开授权地址：", restart=True)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.38)
    add_hyperlink(p, "https://app-data-mcp.luteos.site/auth/metabase/login", "https://app-data-mcp.luteos.site/auth/metabase/login")
    add_number(doc, "输入本人的数据平台账号和密码；密码应与数据平台密码一致。确认信息无误后，点击“确认并授权”。")
    add_number(doc, "身份验证成功后复制页面生成的个人 MCP token。token 形如 appdata_xxxxxx。")
    add_number(doc, "妥善保存 token；不要粘贴到群聊、工单、截图或公开仓库。")
    add_figure(doc, IMG_AUTH, "图 2  app-data MCP 授权页面", width=6.25,
               alt="新版 app-data MCP 授权页面，包含数据平台账号、密码输入框、密码显示开关和确认并授权按钮。")
    add_callout(doc, "授权说明", "个人 MCP token 不按时间自动失效。重新授权会生成新的个人 token，当前账号此前生成或已配置的旧 token 会立即失效；请复制新 token，并及时更新 AI 工具中的 Authorization 配置。", "info")
    add_callout(doc, "地址说明", "授权页面地址为 https://app-data-mcp.luteos.site/auth/metabase/login；WorkBuddy 中配置的 MCP 服务地址为 https://app-data-mcp.luteos.site/mcp。两者使用同一域名，但用途和路径不同，请不要混用。", "warn")

    add_heading(doc, "3.3 第二步：打开 WorkBuddy 自定义连接器", 2)
    add_number(doc, "在 WorkBuddy 左侧导航栏点击“专家·技能·连接器”。", restart=True)
    add_number(doc, "进入页面后切换到顶部“连接器”页签。")
    add_number(doc, "点击右上角“自定义连接器”，打开 MCP 服务管理窗口。")
    add_figure(doc, IMG_WB_CONNECTOR, "图 3  进入 WorkBuddy 自定义连接器", width=6.45,
               alt="WorkBuddy 连接器页面，红色标注依次指向专家技能连接器、连接器页签和自定义连接器入口。")

    add_heading(doc, "3.4 第三步：配置并保存 app-data MCP", 2)
    add_number(doc, "在 MCP 服务管理窗口点击右上角“配置 MCP”。", restart=True)
    add_number(doc, "将配置编辑区中的内容替换为下方 JSON，并把 appdata_xxxxxx 替换为本人最新 token。")
    config = '''{
  "mcpServers": {
    "app-data": {
      "type": "http",
      "url": "https://app-data-mcp.luteos.site/mcp",
      "disabled": false,
      "headers": {
        "Authorization": "Bearer appdata_xxxxxx"
      }
    }
  }
}'''
    add_code_block(doc, config)
    add_number(doc, "确认 JSON 格式完整后点击“保存”；如页面提示未保存，请勿直接关闭窗口。")
    add_figure(doc, IMG_WB_CONFIG, "图 4  粘贴 app-data MCP 配置并保存", width=6.45,
               alt="WorkBuddy MCP 配置编辑窗口，红色标注配置 MCP 按钮、JSON 编辑区和保存按钮。")
    add_callout(doc, "配置检查", "服务地址应为 https://app-data-mcp.luteos.site/mcp；Authorization 必须以 Bearer 加一个空格开头，并使用本人最新 token。", "info")

    add_heading(doc, "3.5 第四步：启用并验证连接器", 2)
    add_number(doc, "保存后返回 MCP 服务列表，确认“我的 MCP”中出现 app-data。", restart=True)
    add_number(doc, "打开 app-data 右侧开关，使其处于绿色启用状态。")
    add_number(doc, "确认 app-data 显示绿色状态点，并显示“10/10 个工具已启用”。")
    add_figure(doc, IMG_WB_ENABLE, "图 5  启用 app-data 并确认工具状态", width=6.2,
               alt="WorkBuddy MCP 服务管理列表，app-data 显示绿色状态点、10/10 个工具已启用，右侧开关为绿色。")
    add_number(doc, "新建对话，发送下方“首次连接验证提示词”，完成一次简单的资产搜索。")
    add_prompt(doc, "首次连接验证提示词（建议直接复制）",
               "请使用 app-data MCP 检查当前连接状态。读取可用工具，确认我的身份与查询权限，并完成一次简单的资产搜索。此步骤不要执行正式数据查询，只需返回：①是否成功连接 app-data MCP；②当前可用的工具类别；③可访问的 Metabase 与 PostHog 资产类型；④本次资产搜索结果。若验证失败，请说明失败环节和处理建议。")
    add_callout(doc, "验证成功的信号", "AI 能确认 app-data MCP 已连接，说明工具与权限状态，并返回实际的资产搜索结果；这只代表接入链路可用，正式分析仍应按第 4 章先确认资产和口径后再查询。", "success")

    add_heading(doc, "3.6 其他 AI 工具", 2)
    add_body(doc, "Codex、Claude Code 等工具的接入逻辑与 WorkBuddy 相同：在其 MCP 配置中新增一个 HTTP 类型服务，填写 MCP URL，并在请求头中配置 Authorization: Bearer <个人 token>。不同工具的配置文件位置、刷新方式和 UI 名称可能不同，请以各工具当前版本说明为准。")
    add_bullet(doc, "服务名称建议统一使用 app-data，便于在提示词中明确指定。")
    add_bullet(doc, "如果工具支持 MCP 工具开关，确认 app-data 已启用。")
    add_bullet(doc, "首次对话仍建议使用本指引提供的连接验证提示词。")
    add_bullet(doc, "不得把个人 token 写入会提交到 Git 的配置文件；优先使用工具的本地私有配置或密钥存储。")

    section4_heading = add_heading(doc, "4. 使用指引：从连接验证到深入分析", 1)
    section4_heading.paragraph_format.page_break_before = True
    add_heading(doc, "4.1 推荐工作流", 2)
    workflow = [
        ("明确问题", "说明要回答的业务问题、时间范围、对象、地区、产品和输出形式。"),
        ("搜索资产", "让 AI 使用 app-data MCP 搜索候选 Model、Metric、Card、Dashboard 或 Insight。"),
        ("确认口径", "核对数据源、字段、筛选条件、去重方式、默认条件和数据时效。"),
        ("执行查询", "在权限允许范围内获取结果；复杂问题可拆为多个可验证的小查询。"),
        ("解释结果", "要求 AI 同时给出数值、口径、资产来源、筛选条件和限制。"),
        ("深入分析", "继续追问趋势、分组、对比、交集、异常和可能原因。"),
    ]
    for idx, (name, detail) in enumerate(workflow):
        p = doc.add_paragraph(style="MCP Number")
        if idx == 0:
            numbering = doc.part.numbering_part.element
            next_id = max([int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))] or [0]) + 1
            num = OxmlElement("w:num")
            num.set(qn("w:numId"), str(next_id))
            abs_id = OxmlElement("w:abstractNumId")
            abs_id.set(qn("w:val"), str(doc._mcp_number_abstract_id))
            num.append(abs_id)
            lvl_override = OxmlElement("w:lvlOverride")
            lvl_override.set(qn("w:ilvl"), "0")
            start_override = OxmlElement("w:startOverride")
            start_override.set(qn("w:val"), "1")
            lvl_override.append(start_override)
            num.append(lvl_override)
            numbering.append(num)
            doc._mcp_current_number_id = next_id
        numPr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        numId = OxmlElement("w:numId")
        numId.set(qn("w:val"), str(doc._mcp_current_number_id))
        numPr.extend([ilvl, numId])
        p._p.get_or_add_pPr().append(numPr)
        set_run_font(p.add_run(f"{name}："), bold=True, color=NAVY)
        set_run_font(p.add_run(detail), color=INK)

    add_heading(doc, "4.2 高质量提问的六个要素", 2)
    rows = [
        ("业务目标", "想回答什么问题", "评估美国地区绑定设备用户规模"),
        ("时间范围", "统计时点或时间段", "截至今天 / 最近 30 天 / 2026 年 6 月"),
        ("分析对象", "用户、设备、事件或订单", "去重 UID、绑定设备用户"),
        ("筛选条件", "地区、产品、状态、人群", "国家=美国，绑定状态=有效"),
        ("分析维度", "总量、趋势、分组、交集", "按周趋势，并按设备型号拆分"),
        ("输出要求", "表格、图表、结论和口径", "给出结果、资产来源、筛选与时效"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [1750, 3000, 4610])
    for i, text in enumerate(("要素", "需要说明", "示例")):
        shade_cell(table.rows[0].cells[i], NAVY)
        set_cell_borders(table.rows[0].cells[i], NAVY)
        set_run_font(table.rows[0].cells[i].paragraphs[0].add_run(text), size=9.5, bold=True, color=WHITE)
    for idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, text in enumerate(row):
            shade_cell(cells[i], LIGHT if idx % 2 else WHITE)
            set_cell_borders(cells[i])
            set_run_font(cells[i].paragraphs[0].add_run(text), size=9.2, bold=(i == 0), color=INK)
    for row in table.rows:
        set_row_cant_split(row)
    add_prompt(doc, "通用提问模板",
               "请使用 app-data MCP，帮我分析【业务目标】。时间范围为【时间范围】，分析对象为【用户/设备/事件】，筛选条件为【地区/产品/状态】，希望按【维度】查看【总量/趋势/分组/交集】。请先搜索并列出候选数据资产，说明口径差异；确认合适资产后再查询。结果中必须包含：数据资产、筛选条件、统计口径、数据时效、核心结果和限制。")

    add_heading(doc, "4.3 从资产搜索开始", 2)
    add_prompt(doc, "资产发现",
               "请使用 app-data MCP 搜索与“设备绑定用户”相关的 Model、Metric、Card 和 Dashboard。列出资产名称、类型、用途、关键字段或指标口径、可用筛选条件，并推荐最适合查询美国地区绑定用户数的资产。先不要执行数据查询。")
    add_prompt(doc, "口径对比",
               "请比较候选资产中“绑定用户数”的统计口径：是否按 UID 去重、绑定状态如何定义、是否限制设备类型、默认时间条件是什么。请把差异列成表格，并说明推荐选择。")

    add_heading(doc, "4.4 单指标查询示例", 2)
    add_prompt(doc, "示例提问",
               "请使用 app-data MCP 查询美国地区绑定设备的去重用户数。优先选择已治理的 Metric，并在查询前确认候选资产的口径。结果中请同时给出：指标名称与资产 ID、底层 Model、country_ad_ch='美国' 的地区筛选、UID 去重方式、内置绑定状态与设备类型条件、数据时效，以及该指标是当前累计总量还是指定时间窗指标。")
    add_figure(doc, IMG_QUERY, "图 6  WorkBuddy 使用 app-data MCP 查询美国地区绑定设备用户数", width=6.35,
               alt="WorkBuddy 查询美国地区绑定设备用户数的示例，返回指标、数据源、地区筛选、统计口径和数据时效。")
    add_callout(doc, "从示例中学习", "截图中的“268,561 人”是一次查询示例，不应脱离查询时间直接复用。一个可复核的回答还应说明 Metric 与底层 Model、country_ad_ch='美国'、UID 去重口径、内置绑定条件和数据时效；若指标统计当前有效绑定用户累计总量，也要明确它不代表某个日期或时间段内的新增、活跃或趋势。", "info")
    page_break(doc)

    add_heading(doc, "4.5 多条件人群交集示例", 2)
    add_prompt(doc, "示例提问",
               "请使用 app-data MCP 查询同时满足以下条件的去重用户：①有有效绑定设备；②社区有正常且未删除的评论；③社区有有效点赞（排除收藏）。请先分别确认三个条件的数据源、用户主键、筛选口径和时间范围，再使用 query_audience 按 uid 计算三条件交集（intersection），不要在客户端拼接。输出各条件人数、交集人数、交集渗透率、查询方式和限制，并生成三集合交集示意图。")
    add_figure(doc, IMG_AUDIENCE_ANALYSIS, "图 7  WorkBuddy 展示三条件用户交集与渗透率", width=6.25,
               alt="WorkBuddy 对有效绑定设备、社区评论和社区点赞三类用户计算交集，展示交集用户数、各集合人数与渗透率。")
    add_callout(doc, "结果解读", "先核对三个 Model 是否都支持 audience 查询，并统一 uid、时间窗和状态过滤。单条件人数之间可能重叠，不能直接相加；交集渗透率要分别说明分母。示例中的交集由 Metabase 内部完成，不是把多份用户明细下载到客户端后合并。", "warn")

    section46_heading = add_heading(doc, "4.6 趋势、分组与行为分析示例", 2)
    section46_heading.paragraph_format.page_break_before = True
    add_heading(doc, "4.6.1 按设备型号查看绑定用户分布", 3)
    add_prompt(doc, "可直接复制的提示词",
               "请使用 app-data MCP 查询各设备型号的绑定用户分布。先确认绑定用户指标、设备型号字段、有效绑定条件和去重口径；按型号输出绑定用户数与占比，并按人数降序排列。请说明各型号人数合计是否可能因同一用户绑定多台设备而大于总体去重用户数，分析头部型号集中度，附数据资产、筛选条件和数据时效，并生成水平条形图。")
    add_figure(doc, IMG_DEVICE_GROUP_ANALYSIS, "图 8  WorkBuddy 分析各设备型号绑定用户分布", width=6.25,
               alt="WorkBuddy 按设备型号分组查询绑定用户数和占比，输出关键发现、数据来源及水平条形图。")
    add_callout(doc, "分组分析检查点", "明确分组字段和占比分母；区分“型号维度下的去重用户数”与“全量去重用户数”。如果一个用户可以绑定多个型号，各组人数可重复计入，因此各组之和不一定等于总体用户数。", "info")

    trend_heading = add_heading(doc, "4.6.2 查看 APP 用户规模与智能设备绑定趋势", 3)
    trend_heading.paragraph_format.page_break_before = True
    add_prompt(doc, "可直接复制的提示词",
               "请使用 app-data MCP 查询最近 15 天 APP 总用户数和智能设备绑定去重用户数趋势。优先选择能在同一日期、同一数据源中提供两个累计指标的预聚合 Model，按天输出 APP 总用户数、智能设备绑定用户数和绑定渗透率（绑定用户数 ÷ APP 总用户数）。请说明两个指标的用户主键、绑定状态、智能设备范围、累计或新增口径、日期字段和数据时效；分别计算期初到期末的增量、增长率及渗透率变化（百分点），标记异常日期，并生成双轴折线图：左轴展示 APP 总用户数，右轴展示智能设备绑定用户数。")
    add_figure(doc, IMG_APP_BIND_TREND_ANALYSIS, "图 9  WorkBuddy 查询 APP 用户规模与智能设备绑定趋势", width=6.25,
               alt="WorkBuddy 使用同源预聚合模型按日查询 APP 总用户数、智能设备绑定用户数和绑定渗透率。")
    add_figure(doc, IMG_APP_BIND_DUAL_AXIS_TREND, "图 10  WorkBuddy 展示 APP 用户与智能设备绑定用户双轴趋势", width=6.25,
               alt="WorkBuddy 使用双轴折线图展示 APP 总用户累计规模和智能设备绑定用户累计规模，左右纵轴分别对应两个指标。")
    add_callout(doc, "趋势分析检查点", "两个指标必须使用同一用户范围、同一日期和兼容的去重口径，渗透率才能直接计算。双轴图的左右刻度不同，两条线视觉上接近不代表增量或增长率一致；应结合每日增量、增长率和渗透率百分点变化解读。截图数值仅为当次查询示例，也不要把累计规模上涨误读为活跃度提升。", "warn")

    posthog_heading = add_heading(doc, "4.6.3 PostHog 行为分析与看板导航", 3)
    posthog_heading.paragraph_format.page_break_before = True
    add_body(doc, "PostHog 的 Dashboard 适合定位一个业务主题下的分析入口，Insight 则承载具体的趋势、漏斗、留存或分维度查询。推荐先盘点看板及关联洞察，再选择与问题匹配的 Insight 获取实时结果，最后把多项结果汇总成可复核的分析报告。")
    add_prompt(doc, "第一步：盘点看板与关联洞察",
               "请使用 app-data MCP 搜索 PostHog 中与【业务主题】相关的 Dashboard 和 Insight。先列出匹配的看板名称、项目、关联 Insight 数量；再逐项列出 Insight 名称、类型（趋势/漏斗/留存等）、统计时间范围和回答的业务问题。此步骤先完成资产导航与口径盘点，不要把不同时间窗或不同计算类型的结果直接比较。")
    add_figure(doc, IMG_POSTHOG_INSIGHT_INVENTORY, "图 11  WorkBuddy 盘点 PostHog 看板及关联洞察", width=6.25,
               alt="WorkBuddy 使用 app-data MCP 查询 PostHog 小工具相关看板和洞察，列出洞察名称、分析类型与时间范围。")
    add_callout(doc, "资产导航检查点", "确认 PostHog project、Dashboard 名称及关联 Insight 是否完整。趋势、分维度趋势、有序漏斗和无序漏斗回答的问题不同；即使名称相近，也要核对事件顺序、去重方式、属性过滤与时间范围后再解释。", "info")

    add_prompt(doc, "第二步：查询洞察并生成分析报告",
               "请继续使用 app-data MCP 查询上一步确认的 PostHog Insight 实时数据。围绕【核心目标】汇总转化率、渗透率、页面 PV/UV、入口点击或其他关键指标，统一标注各指标的事件定义、属性筛选、计算方式和时间范围；识别趋势变化、头部贡献和异常点，把“查询事实”“基于数据的推断”“待验证假设”分开表达。请生成包含指标概览、趋势图、明细表和关键发现的可视化报告，并保留 Dashboard/Insight 来源，方便返回原看板继续核验。")
    add_figure(doc, IMG_POSTHOG_ANALYSIS_REPORT, "图 12  WorkBuddy 汇总 PostHog 洞察并生成可视化分析报告", width=6.25,
               alt="WorkBuddy 汇总 PostHog 小工具洞察的关键发现，并在右侧打开包含指标卡片和趋势图的可视化分析报告。")
    add_callout(doc, "报告解读检查点", "截图展示的是一次“小工具”分析示例。报告中的转化率、渗透率、PV/UV 和趋势可能来自不同 Insight 与时间窗，必须分别保留来源和口径；AI 生成的报告用于汇总与导航，正式结论仍应回到 PostHog 原始 Dashboard/Insight 核验。", "warn")

    add_heading(doc, "4.7 如何继续追问做深入分析", 2)
    add_body(doc, "拿到第一个结果后，不要立即把数值当成结论。建议沿着“确认 → 对比 → 拆解 → 归因 → 验证”逐步追问。")
    for text in (
        "确认：这个结果使用了哪个资产？指标定义、时间范围、去重口径和默认过滤是什么？",
        "对比：与上一周期、目标值、其他地区或其他型号相比如何？差异是多少？",
        "拆解：按地区、渠道、产品、设备型号、新老用户或关键行为拆分，差异由哪部分贡献？",
        "归因：有哪些数据事实支持可能原因？哪些只是推测？还需要查询什么才能验证？",
        "验证：换用另一个已治理资产交叉验证，或检查 Dashboard/Insight 中是否有一致趋势。",
    ):
        add_bullet(doc, text)
    add_prompt(doc, "推荐的分析约束",
               "请把“查询事实”“基于数据的推断”“尚待验证的假设”分开表达。所有结论必须附数据资产、筛选条件、时间范围和口径；如果现有资产不足以回答，请明确说明缺口，不要自行编造数据。")

    add_heading(doc, "5. 安全规范、常见问题与检查清单", 1)
    add_heading(doc, "5.1 安全与权限规范", 2)
    add_bullet(doc, "账号、密码和个人 token 仅限本人使用，不得共享、转发或写入公开文档。")
    add_bullet(doc, "不要把真实 token 提交到 Git、粘贴到群聊、截图或问题工单；文档和示例始终使用 appdata_xxxxxx。")
    add_bullet(doc, "查询能力受个人数据权限控制。遇到无权限资产，应走正常权限申请流程。")
    add_bullet(doc, "结果仅用于授权范围内的工作分析；涉及敏感用户数据时遵守公司数据安全和最小必要原则。")
    add_bullet(doc, "AI 输出可能存在解释偏差。重要经营结论、对外数据和关键决策须复核资产、口径与原始结果。")

    add_heading(doc, "5.2 常见问题", 2)
    faq = [
        ("AI 没有识别到 app-data MCP", "检查 JSON 格式、服务名、URL、disabled=false 和 Authorization 请求头；保存后重启或刷新工具，并重新发送首次连接验证提示词。"),
        ("返回 401 / 未认证", "确认 Authorization 的格式为 Bearer + 空格 + 个人 token；不要包含引号外的多余空格。必要时重新完成授权。"),
        ("返回 403 / 无权限", "当前账号没有目标资产或数据的访问权限。请确认使用的 Metabase 账号，并按正常流程申请权限。"),
        ("能搜索资产但查询失败", "先让 AI 展示资产类型、所需参数、字段与过滤条件；缩小时间范围，减少维度，并拆分复杂查询。"),
        ("结果与预期不一致", "核对资产、时间窗、地区字段、去重主键、默认过滤、状态条件和数据时效；必要时用第二个已治理资产交叉验证。"),
        ("查询很慢或超时", "缩短时间范围、减少分组维度、先查汇总后查明细；将大问题拆成多个小查询。"),
        ("AI 只给分析建议，没有调用数据", "明确写出“请使用 app-data MCP”，要求先搜索资产并在答案中列出实际使用的资产、筛选条件、口径与时效。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [3000, 6360])
    for i, text in enumerate(("问题", "处理建议")):
        shade_cell(table.rows[0].cells[i], NAVY)
        set_cell_borders(table.rows[0].cells[i], NAVY)
        set_run_font(table.rows[0].cells[i].paragraphs[0].add_run(text), size=9.5, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for idx, row in enumerate(faq):
        cells = table.add_row().cells
        for i, text in enumerate(row):
            shade_cell(cells[i], LIGHT if idx % 2 else WHITE)
            set_cell_borders(cells[i])
            set_run_font(cells[i].paragraphs[0].add_run(text), size=9.0, bold=(i == 0), color=INK)
    for row in table.rows:
        set_row_cant_split(row)
    add_heading(doc, "5.3 五分钟接入检查清单", 2)
    checks = [
        "已使用本人 Metabase 账号完成 app-data MCP 授权。",
        "已复制个人 token，并仅保存在本人设备的私有配置中。",
        "WorkBuddy 的 app-data URL、disabled 和 Authorization 配置正确。",
        "已重启或刷新 WorkBuddy，app-data MCP 处于启用状态。",
        "已发送首次连接验证提示词，AI 能确认连接、工具与权限状态，并返回资产搜索结果。",
        "已完成一次资产搜索，并核对名称、类型和业务口径。",
        "已完成一次简单查询，回答包含资产、筛选、口径和数据时效。",
        "知道如何通过趋势、分组、对比和交集继续追问。",
        "知道个人 token、账号密码和敏感数据不可分享。",
    ]
    for item in checks:
        p = doc.add_paragraph(style="MCP Bullet")
        set_run_font(p.add_run("□  "), size=12, bold=True, color=BLUE)
        set_run_font(p.add_run(item), color=INK)

    add_heading(doc, "附录：日常数据分析开场模板", 2)
    add_prompt(doc, "开始一次正式数据分析",
               "请使用 app-data MCP 分析【我的业务问题】。第一步，搜索相关的 Metabase Model/Metric/Card/Dashboard 和 PostHog Insight/Dashboard；第二步，列出候选资产的名称、用途、口径差异和适用场景，等待我确认后再查询；第三步，按确认的资产执行查询，返回数据资产、筛选条件、统计口径、数据时效、核心结果和限制；第四步，把查询事实、基于数据的推断和待验证假设分开表达。如现有资产不足以回答，请明确说明缺口，不要猜测字段或编造数据。现在先完成前两步。")
    add_callout(doc, "分析完成标准", "在确认资产与口径后返回可复核数据，清楚说明筛选、时效和限制，并能基于查询事实支持后续对比、拆解与验证。", "success")

    # Core properties intentionally generic for internal sharing.
    doc.core_properties.title = "app-data-mcp架构介绍、配置接入及使用指引"
    doc.core_properties.subject = "APP生态中心 app-data-mcp 内部接入与使用手册"
    doc.core_properties.author = "APP生态中心数据平台"
    doc.core_properties.last_modified_by = "APP生态中心数据平台"
    doc.core_properties.keywords = "app-data-mcp, MCP, Metabase, PostHog, WorkBuddy"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
